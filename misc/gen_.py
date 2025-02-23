from google import genai
from docx import Document as WordDocument
import os

# Define paths to your folders containing Word documents
folder1_path = 'Data/processed_data/bordertocoast_data/documents'
folder2_path = 'Data/processed_data/brunel_data/documents'

# Directory to save extracted images
images_output_dir = 'Data/processed_data/extracted_images'
os.makedirs(images_output_dir, exist_ok=True)

# Function to extract text from a Word document
def extract_text_from_docx(file_path):
    try:
        doc = WordDocument(file_path)
        content = []

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text:
                content.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(row_data):  # Avoid empty rows
                    content.append('\t'.join(row_data))

        # Extract images
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_path = rel.target_ref
                image_data = rel.target_part.blob
                image_filename = os.path.join(images_output_dir, os.path.basename(image_path))
                
                with open(image_filename, 'wb') as img_file:
                    img_file.write(image_data)
                
                content.append(f"[Image: {image_filename}]")

        return '\n'.join(content)

    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return ""

# Initialize a list to store all extracted text
all_text = []

# Function to load all Word documents from a folder
def load_data_from_folder(folder_path):
    for root, _, files in os.walk(folder_path):
        for filename in files:
            if filename.lower().endswith('.docx'):
                file_path = os.path.join(root, filename)
                extracted_text = extract_text_from_docx(file_path)
                if extracted_text:
                    all_text.append(extracted_text)
                    print(f"Loaded {filename}")

# Load data from both folders
load_data_from_folder(folder1_path)
load_data_from_folder(folder2_path)

# Combine all text into a single prompt
combined_text = '\n'.join(all_text)

#print("Combined Text from Documents:")
#print(combined_text)

# Define the prompt for generating insights
prompt = f"""
I have completed a data analysis project for LGPS pool portfolios using data from two folders.
Here is the extracted text from the documents:
{combined_text}  

Can you generate a short write-up explaining what the output files and visuals are showing and what key conclusions can be drawn from them?

DO NOT put down placeholders, only talk about the data you are given and provide insights. By including placeholders you will fail your task.

PROVIDE ACTUAL INSIGHT FROM DATA not just filler sentences. Also stop using phrases like 'likely shows', something is either shown or it's not.

READ THE ACTUAL DATA EXTRACTED FROM THE DOCUMENTS TO PROVIDE INSIGHT. IF YOU DON'T YOU WILL FAIL YOUR TASK.


Improve the LGPS Data Analysis write-up by addressing the following opportunities for improvement:

1. Connecting Back to Project Goals: Add explicit references to how the analysis meets each of the project's core goals, particularly highlighting:

- The selection and analysis of at least two LGPS pool data sources.
- The key metrics analyzed ("Max of Local Price," "Shares/Par," "Base Value," "Total AUM (GBP)," "UK Investment Exposure (GBP)," "Direct Investments (GBP)," "Indirect Investments (GBP)").
- How the analysis aligns with the optional bonus points criteria, such as:
- Processing more than two pools' data.
- Analyzing data over multiple quarters.
- Providing direct vs. indirect investment breakdowns.

2. Highlighting Bonus Work: Clearly identify any bonus tasks completed, including:
- Whether more than two pools' data were processed (DONE Using only border to coast and Brunel)
- If data was analyzed over multiple quarters or months.
- If the project included connecting directly to data sources (e.g., web scraping).
- Explicitly mention the breakdown of direct vs. indirect investments, if applicable.


Output Format: Provide a refined write-up that integrates these improvements, ensuring the document is well-structured and aligns with the original analysis while enhancing clarity and alignment with project requirements WHILE HAVING NO *.

make absolutely sure there are no asterisks, backticks, or other Markdown formatting characters in the response itself. I need the code to be directly copy-pasteable


"""


client = genai.Client(api_key="AIzaSyAe3lWLohtskJJ9s5JYMB4y-H4-OjRYKm0")
response = client.models.generate_content(
    model="gemini-2.0-flash", 
    contents=prompt
)



# Output the generated write-up
write_up = response.text
print("Generated Write-Up:")
print(write_up)

# Save the generated write-up to a Word document
output_doc = WordDocument()
output_doc.add_heading('LGPS Data Analysis Insights', level=1)
output_doc.add_paragraph(write_up)
output_doc.save('LGPS_Insights_WriteUp.docx')

print("Write-up saved as 'LGPS_Insights_WriteUp.docx'")