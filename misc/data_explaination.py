import os
from google import genai
from docx import Document as WordDocument

class LGPSDataAnalyzer:
    def __init__(self, api_key: str, folder1_path: str, folder2_path: str, images_output_dir: str, model_name: str = "gemini-2.0-flash"):
        self.client = genai.Client(api_key=api_key)
        self.folder1_path = folder1_path
        self.folder2_path = folder2_path
        self.images_output_dir = images_output_dir
        self.model_name = model_name
        os.makedirs(self.images_output_dir, exist_ok=True)
        self.all_text = []

    def extract_text_from_docx(self, file_path: str) -> str:
        try:
            doc = WordDocument(file_path)
            content = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    content.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data):
                        content.append('\t'.join(row_data))

            # Extract images
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_path = rel.target_ref
                    image_data = rel.target_part.blob
                    image_filename = os.path.join(self.images_output_dir, os.path.basename(image_path))
                    with open(image_filename, 'wb') as img_file:
                        img_file.write(image_data)
                    content.append(f"[Image: {image_filename}]")

            return '\n'.join(content)

        except Exception as e:
            print(f"Error reading {file_path}: {e}")
            return ""

    def load_data_from_folder(self, folder_path: str) -> None:
        for root, _, files in os.walk(folder_path):
            for filename in files:
                if filename.lower().endswith('.docx'):
                    file_path = os.path.join(root, filename)
                    extracted_text = self.extract_text_from_docx(file_path)
                    if extracted_text:
                        self.all_text.append(extracted_text)
                        print(f"Loaded {filename}")

    def generate_insights(self, combined_text: str, prompt_template: str) -> str:
        full_prompt = prompt_template.format(combined_text=combined_text)
        response = self.client.models.generate_content(
            model=self.model_name,
            contents=full_prompt
        )
        return response.text

    def save_write_up(self, write_up: str, output_filename: str = 'LGPS_Insights_WriteUp.docx') -> None:
        output_doc = WordDocument()
        output_doc.add_heading('LGPS Data Analysis Insights', level=1)
        output_doc.add_paragraph(write_up)
        output_doc.save(output_filename)
        print(f"Write-up saved as '{output_filename}'")

    def run_analysis(self, prompt_template: str) -> None:
        self.load_data_from_folder(self.folder1_path)
        self.load_data_from_folder(self.folder2_path)
        combined_text = '\n'.join(self.all_text)
        write_up = self.generate_insights(combined_text, prompt_template)
        #print("Generated Write-Up:")
        #print(write_up)
        self.save_write_up(write_up)

# Run Program 
def run(): 
    api_key = "AIzaSyAe3lWLohtskJJ9s5JYMB4y-H4-OjRYKm0"
    folder1_path = 'Data/processed_data/bordertocoast_data/documents'
    folder2_path = 'Data/processed_data/brunel_data/documents'
    images_output_dir = 'Data/processed_data/extracted_images'

    prompt_template = '''
    I have completed a data analysis project for LGPS pool portfolios using data from two folders.
    Here is the extracted text from the documents:
    {combined_text}  
    Can you generate a short write-up explaining what the output files and visuals are showing and what key conclusions can be drawn from them?
    DO NOT put down placeholders, only talk about the data you are given and provide insights. By including placeholders you will fail your task.
    ...
    '''
    analyzer = LGPSDataAnalyzer(api_key, folder1_path, folder2_path, images_output_dir)
    analyzer.run_analysis(prompt_template)

# Example usage
if __name__ == '__main__':
    run()
