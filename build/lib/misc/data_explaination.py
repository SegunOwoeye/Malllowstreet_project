import os
import google.generativeai as genai
from docx import Document as WordDocument

# Set your Google Gemini API key
genai.configure(api_key="YOUR_API_KEY")

# Define paths to your folders containing Word documents
folder1_path = 'Data/processed_data/bordertocoast_data/documents'
folder2_path = 'Data/processed_data/brunel_data/documents'

# Function to extract text from a Word document
def extract_text_from_docx(file_path):
    try:
        doc = WordDocument(file_path)
        text = [paragraph.text for paragraph in doc.paragraphs if paragraph.text]
        return '\n'.join(text)
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return ""

# Initialize a list to store all extracted text
all_text = []

# Function to load all Word documents from a folder
def load_data_from_folder(folder_path):
    for filename in os.listdir(folder_path):
        if filename.endswith('.docx'):
            file_path = os.path.join(folder_path, filename)
            extracted_text = extract_text_from_docx(file_path)
            if extracted_text:
                all_text.append(extracted_text)
                print(f"Loaded {filename}")

# Load data from both folders
load_data_from_folder(folder1_path)
load_data_from_folder(folder2_path)

# Combine all text into a single prompt
combined_text = '\n'.join(all_text)

# Define the prompt for generating insights
prompt = f"""
I have completed a data analysis project for LGPS pool portfolios using data from two folders.
Here is the extracted text from the documents:
{combined_text[:4000]}  # Limit to 4000 characters to fit input limits

Can you generate a short write-up explaining what the output files and visuals are showing and what key conclusions can be drawn from them?
"""

# Define generation configuration
generation_config = {
    "temperature": 0.7,
    "top_p": 0.95,
    "top_k": 64,
    "max_output_tokens": 8192,
    "response_mime_type": "text/plain",
}

# Create the generative model for Google Gemini
model = genai.GenerativeModel(
    model_name="gemini-1.5-flash",
    generation_config=generation_config,
    system_instruction="Act as a financial analyst providing insights into LGPS pool portfolio data analysis."
)

# Generate the response using Google Gemini API
response = model.generate_content(prompt)

# Output the generated write-up
write_up = response.text
print(write_up)

# Save the generated write-up to a Word document
output_doc = WordDocument()
output_doc.add_heading('LGPS Data Analysis Insights', level=1)
output_doc.add_paragraph(write_up)
output_doc.save('LGPS_Insights_WriteUp.docx')

print("Write-up saved as 'LGPS_Insights_WriteUp.docx'")
