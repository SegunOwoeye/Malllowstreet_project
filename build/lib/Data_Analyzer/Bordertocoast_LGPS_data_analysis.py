import os
import re
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import PyPDF2

class LGPSDataAnalysis:
    def __init__(self, input_directory: str, output_directory: str, images_directory: str):
        self.input_directory = input_directory
        self.output_directory = output_directory
        self.images_directory = images_directory
        os.makedirs(self.output_directory, exist_ok=True)
        os.makedirs(self.images_directory, exist_ok=True)

    def extract_text_from_pdf(self, pdf_path: str) -> str:
        pdf_text = ""
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                pdf_text += page.extract_text() if page.extract_text() else ""
        return pdf_text

    def extract_numerical_data(self, pdf_text: str) -> pd.DataFrame:
        pattern = r"(\d{1,3}(?:,\d{3})*(?:\.\d{2})?)\s+(\d{1,3}(?:,\d{3})*)\s+(\d{1,3}(?:,\d{3})*(?:\.\d{2})?)"
        matches = re.findall(pattern, pdf_text)
        numerical_columns = ['Max of Local Price', 'Shares/Par', 'Base Value']
        numerical_data = pd.DataFrame(matches, columns=numerical_columns)
        numerical_data['Max of Local Price'] = pd.to_numeric(numerical_data['Max of Local Price'].str.replace(',', ''), errors='coerce')
        numerical_data['Shares/Par'] = pd.to_numeric(numerical_data['Shares/Par'].str.replace(',', ''), errors='coerce')
        numerical_data['Base Value'] = pd.to_numeric(numerical_data['Base Value'].str.replace(',', ''), errors='coerce')
        return numerical_data

    def generate_report(self, filename: str, numerical_data: pd.DataFrame):
        numerical_summary = numerical_data.describe().T
        numerical_summary['Sum'] = numerical_data.sum()
        numerical_summary['Median'] = numerical_data.median()
        numerical_summary['Standard Deviation'] = numerical_data.std()
        numerical_summary['Variance'] = numerical_data.var()
        doc = Document()
        doc.add_heading(f'LGPS Data Analysis - {filename}', level=1)
        doc.add_heading('Detailed Numerical Analysis', level=2)
        doc.add_paragraph('The following table provides detailed statistics of the numerical data:')
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Value'
        for index, row in numerical_summary.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(index)
            row_cells[1].text = f"{row['Sum']:.2f}"

        # Generate and save the histogram for 'Max of Local Price'
        plt.figure(figsize=(8, 6))
        plt.hist(numerical_data['Max of Local Price'], bins=50, color='blue', alpha=0.7)
        plt.xlabel('Max of Local Price')
        plt.ylabel('Frequency')
        plt.title('Distribution of Max of Local Price')
        histogram_path = os.path.join(self.images_directory, f'{os.path.splitext(filename)[0]}_histogram.png')
        plt.savefig(histogram_path)
        plt.close()
        doc.add_picture(histogram_path, width=Inches(6))

        # Generate and save the scatter plot for 'Shares/Par' vs 'Base Value'
        plt.figure(figsize=(8, 6))
        plt.scatter(numerical_data['Shares/Par'], numerical_data['Base Value'], alpha=0.6, color='green')
        plt.xlabel('Shares/Par')
        plt.ylabel('Base Value')
        plt.title('Shares/Par vs Base Value')
        scatter_plot_path = os.path.join(self.images_directory, f'{os.path.splitext(filename)[0]}_scatter.png')
        plt.savefig(scatter_plot_path)
        plt.close()
        doc.add_picture(scatter_plot_path, width=Inches(6))

        # Save the document to the output directory
        output_file_path = os.path.join(self.output_directory, f'{os.path.splitext(filename)[0]}_Analysis_Report.docx')
        doc.save(output_file_path)
        print(f"Analysis report generated: {output_file_path}")

    def process_all_pdfs(self):
        for filename in os.listdir(self.input_directory):
            if filename.lower().endswith('.pdf'):
                pdf_path = os.path.join(self.input_directory, filename)
                pdf_text = self.extract_text_from_pdf(pdf_path)
                numerical_data = self.extract_numerical_data(pdf_text)
                self.generate_report(filename, numerical_data)
        print("Processing complete!")


# Runs the core program functions
def run():
    input_directory = 'Data/raw_data/bordertocoast_data/pdf'
    output_directory = 'Data/processed_data/bordertocoast_data/documents/'
    images_directory = 'Data/processed_data/bordertocoast_data/supporting_files/'
    analyzer = LGPSDataAnalysis(input_directory, output_directory, images_directory)
    analyzer.process_all_pdfs()


if __name__ == '__main__':
    run()