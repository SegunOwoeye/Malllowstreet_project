import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import os

class LGPSDataAnalyzer:
    def __init__(self, input_directory, output_directory, supporting_files_directory):
        self.input_directory = input_directory
        self.output_directory = output_directory
        self.supporting_files_directory = supporting_files_directory
        os.makedirs(self.output_directory, exist_ok=True)
        os.makedirs(self.supporting_files_directory, exist_ok=True)

    def get_excel_files_and_sheets(self):
        file_sheet_map = {}
        for file_name in os.listdir(self.input_directory):
            if file_name.endswith('.xlsx'):
                file_path = os.path.join(self.input_directory, file_name)
                try:
                    excel_data = pd.ExcelFile(file_path)
                    file_sheet_map[file_path] = excel_data.sheet_names
                except Exception as e:
                    print(f'Error reading {file_name}: {e}')
        return file_sheet_map

    def load_data(self, file_path, sheet_name):
        data = pd.read_excel(file_path, sheet_name=sheet_name)
        data.columns = data.columns.str.strip()
        return data

    def analyze_data(self, data):
        total_aum = data['Base Market Value'].sum()
        uk_investment = data[
            (data['Issue Country Name'] == 'UNITED KINGDOM') |
            (data['Incorporated Country Name'] == 'UNITED KINGDOM')
        ]['Base Market Value'].sum()
        asset_class_breakdown = data.groupby('Investment Type Name')['Base Market Value'].sum().reset_index()
        sector_breakdown = data.groupby('Major Industry Name')['Base Market Value'].sum().reset_index()
        direct_investments = data[
            ~data['Investment Type Name'].str.contains('FUND', na=False, case=False)
        ]['Base Market Value'].sum()
        indirect_investments = data[
            data['Investment Type Name'].str.contains('FUND', na=False, case=False)
        ]['Base Market Value'].sum()
        summary_data = pd.DataFrame({
            'Metric': [
                'Total AUM (GBP)',
                'UK Investment Exposure (GBP)',
                'Direct Investments (GBP)',
                'Indirect Investments (GBP)'
            ],
            'Value': [
                total_aum,
                uk_investment,
                direct_investments,
                indirect_investments
            ]
        })
        return summary_data, asset_class_breakdown, sector_breakdown

    def add_table_to_doc(self, doc, df, title):
        doc.add_heading(title, level=2)
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.autofit = True
        hdr_cells = table.rows[0].cells
        for i, column_name in enumerate(df.columns):
            hdr_cells[i].text = str(column_name)
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)

    def generate_report(self, file_path, sheet_name, summary_data, asset_class_breakdown, sector_breakdown):
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        doc = Document()
        doc.add_heading(f'{base_name} - LGPS Data Analysis Report', level=1)
        self.add_table_to_doc(doc, summary_data, 'Summary Metrics')
        self.add_table_to_doc(doc, asset_class_breakdown, 'Asset Class Breakdown')
        self.add_table_to_doc(doc, sector_breakdown, 'Sector Breakdown')

        asset_class_chart_path = os.path.join(self.supporting_files_directory, f'{base_name}_asset_class_breakdown.png')
        plt.figure(figsize=(8, 6))
        plt.pie(asset_class_breakdown['Base Market Value'], labels=asset_class_breakdown['Investment Type Name'], autopct='%1.1f%%')
        plt.title('Asset Class Breakdown')
        plt.savefig(asset_class_chart_path)
        plt.close()
        doc.add_picture(asset_class_chart_path, width=Inches(5.5))

        sector_chart_path = os.path.join(self.supporting_files_directory, f'{base_name}_sector_breakdown.png')
        plt.figure(figsize=(10, 6))
        plt.bar(sector_breakdown['Major Industry Name'], sector_breakdown['Base Market Value'])
        plt.xticks(rotation=45, ha='right')
        plt.title('Sector Breakdown')
        plt.xlabel('Sector')
        plt.ylabel('Base Market Value (GBP)')
        plt.tight_layout()
        plt.savefig(sector_chart_path)
        plt.close()
        doc.add_picture(sector_chart_path, width=Inches(5.5))

        output_path = os.path.join(self.output_directory, f'{base_name}_LGPS_Data_Analysis_Report.docx')
        doc.save(output_path)
        print(f'Report generated: {output_path}')

    def run_analysis(self):
        file_sheet_map = self.get_excel_files_and_sheets()
        for file_path, sheet_names in file_sheet_map.items():
            for sheet_name in sheet_names:
                print(f'Processing {file_path} - {sheet_name}')
                data = self.load_data(file_path, sheet_name)
                summary_data, asset_class_breakdown, sector_breakdown = self.analyze_data(data)
                self.generate_report(file_path, sheet_name, summary_data, asset_class_breakdown, sector_breakdown)

# Runs Main Class
def run():
    input_dir = 'Data/raw_data/brunel_data/'
    output_dir = 'Data/processed_data/brunel_data/documents/'
    supporting_files_dir = 'Data/processed_data/brunel_data/supporting_files/'
    analyzer = LGPSDataAnalyzer(input_dir, output_dir, supporting_files_dir)
    analyzer.run_analysis()

# Example usage
if __name__ == '__main__':
    run()
