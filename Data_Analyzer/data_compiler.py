from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Inches

def extract_fund_name_and_market(filename):
    """Extracts the fund name and market from the filename."""
    parts = filename.replace("_Analysis_Report.docx", "").split("-")
    fund_name = "-".join(parts[:-2])
    market = parts[-2]
    return fund_name.strip(), market.strip()

def extract_table_from_docx(file_path):
    """Extracts the numerical data table from a Word document."""
    doc = Document(file_path)
    for table in doc.tables:
        data = []
        for row in table.rows:
            data.append([cell.text.strip() for cell in row.cells])
        df = pd.DataFrame(data)
        if df.shape[1] == 2:  # Ensures it's a metric-value table
            return df
    return None

def compile_financial_data(input_folder, output_file):
    """Compiles financial data from multiple Word reports into a structured summary document."""
    input_path = Path(input_folder)
    output_path = Path(output_file)
    
    # Ensure the output directory exists
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    compiled_data = []
    
    if not any(input_path.iterdir()):
        print(f"No Word documents found in {input_folder}")
        return
    
    for file in input_path.glob("*.docx"):
        fund_name, market = extract_fund_name_and_market(file.name)
        df = extract_table_from_docx(file)
        
        if df is not None:
            df.columns = ["Metric", "Value"]
            df.insert(0, "Market", market)
            df.insert(0, "Fund Name", fund_name)
            compiled_data.append(df)
        else:
            print(f"No valid data table found in {file.name}, skipping.")
    
    # Merge all extracted data
    if not compiled_data:
        print("No valid data extracted from documents.")
        return
    final_df = pd.concat(compiled_data, ignore_index=True)
    
    # Create a new Word document
    output_doc = Document()
    output_doc.add_heading('LGPS Compiled Financial Report', level=1)
    
    table = output_doc.add_table(rows=1, cols=len(final_df.columns))
    table.autofit = True
    
    hdr_cells = table.rows[0].cells
    for i, column_name in enumerate(final_df.columns):
        hdr_cells[i].text = str(column_name)
    
    for _, row in final_df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
    
    # Save the output document
    output_doc.save(output_path)
    print(f"Compiled report saved as: {output_file}")

# Run the script
if __name__ == "__main__":
    input_folder = "Data/processed_data/brunel_data/documents/"  # Adjusted path
    output_file = "Data/processed_data/brunel_data/reports/supporting_documents/LGPS_Compiled_Financial_Report.docx"
    compile_financial_data(input_folder, output_file)