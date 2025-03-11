import os
import re
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# ---------------------------
# Step 1: Extract lines from DOCX
# ---------------------------
def extract_lines_from_docx(file_path):
    doc = Document(file_path)
    lines = []
    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    # Also extract text from tables (if any)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    lines.append(cell_text)
    return lines

# ---------------------------
# Step 2: Filter out repeated header labels
# ---------------------------
# These labels appear in the document as headers – we remove them.
HEADER_LABELS = {"LGPS Compiled Financial Report", "Fund Name", "Market", "Metric", "Value"}

def filter_lines(lines):
    return [line for line in lines if line not in HEADER_LABELS]

# ---------------------------
# Step 3: Parse the lines into records using a state machine
# ---------------------------
def is_fund_header(line):
    # A simple heuristic: if the line contains a hyphen and either 'Report' or 'Fund'
    return '-' in line and (("Report" in line) or ("Fund" in line))

def is_market(line):
    # If the line is a year (digits) or equals "Fund" (used in some cases)
    return line.isdigit() or line.lower() == "fund"

def parse_records(lines):
    records = []
    current_fund = None
    current_market = None
    i = 0
    while i < len(lines):
        # If the current line looks like a header and (either we have no current fund or it's different)
        if is_fund_header(lines[i]) and (current_fund is None or lines[i] != current_fund):
            # Expect next line to be market info if available
            if i + 1 < len(lines) and is_market(lines[i+1]):
                current_fund = lines[i]
                current_market = lines[i+1]
                i += 2
                continue
        # If the current line is a duplicate header (same as current_fund), skip it with its following market info
        if current_fund is not None and lines[i] == current_fund:
            if i + 1 < len(lines) and is_market(lines[i+1]):
                i += 2
                continue
        # Otherwise, assume the next two lines are metric and value (if available)
        if i + 1 < len(lines):
            metric = lines[i]
            value = lines[i+1]
            # Only add a record if we already have a current fund header set
            if current_fund is not None:
                records.append({
                    "Fund": current_fund,
                    "Market": current_market,
                    "Metric": metric,
                    "Value": value
                })
            i += 2
        else:
            break
    return records

# ---------------------------
# Step 4: Build DataFrame and pivot
# ---------------------------
def build_pivot_dataframe(records):
    df = pd.DataFrame(records)
    # Convert the Value column to numeric, removing commas and coercing errors
    df['Value'] = pd.to_numeric(df['Value'].str.replace(',', ''), errors='coerce')
    # Use pivot_table with an aggregation function to handle duplicates (using 'first' here)
    pivot_df = df.pivot_table(index="Fund", columns="Metric", values="Value", aggfunc='first').reset_index()
    return pivot_df


# ---------------------------
# Step 5: Create a graph (e.g., bar chart for Base Value)
# ---------------------------
def create_base_value_chart(pivot_df, output_image="base_value_bar_chart.png"):
    if "Base Value" not in pivot_df.columns:
        print("No 'Base Value' metric found. Skipping graph creation.")
        return None
    plt.figure(figsize=(8, 6))
    plt.bar(pivot_df['Fund'], pivot_df['Base Value'])
    plt.xlabel('Fund')
    plt.ylabel('Base Value')
    plt.title('Base Value by Fund')
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()
    plt.savefig(output_image)
    plt.close()
    return output_image

# ---------------------------
# Step 6: Generate a new DOCX report with table and graph
# ---------------------------
def add_table_to_doc(doc, df, title):
    doc.add_heading(title, level=2)
    # Create table with a header row
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.autofit = True
    hdr_cells = table.rows[0].cells
    for idx, col_name in enumerate(df.columns):
        hdr_cells[idx].text = str(col_name)
    # Add rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for idx, value in enumerate(row):
            row_cells[idx].text = str(value)
    doc.add_paragraph("")  # add spacing

def generate_report(pivot_df, graph_path, output_docx="LGPS_Compiled_Financial_Report_Organized.docx"):
    doc = Document()
    doc.add_heading('LGPS Compiled Financial Report - Organized', level=1)
    doc.add_paragraph("This report organizes the raw LGPS data into a structured table and provides a graphical analysis.")
    
    # Add the pivot table
    add_table_to_doc(doc, pivot_df, "Summary Table")
    
    # Add the graph if available
    if graph_path and os.path.exists(graph_path):
        doc.add_heading("Graphical Analysis", level=2)
        doc.add_picture(graph_path, width=Inches(6))
    
    doc.save(output_docx)
    print(f"Report generated and saved as '{output_docx}'")

# ---------------------------
# Main function
# ---------------------------
def main(input_path, output_path):
    input_docx = f"{input_path}LGPS_Compiled_Financial_Report.docx"  # your input file
    output_docx = f"{output_path}LGPS_Compiled_Financial_Report_Organized.docx"
    graph_image = f"{output_path}base_value_bar_chart.png"
    
    # 1. Extract and filter lines from the DOCX
    raw_lines = extract_lines_from_docx(input_docx)
    filtered_lines = filter_lines(raw_lines)
    
    # 2. Parse records using header detection and metric-value pairing
    records = parse_records(filtered_lines)
    if not records:
        print("No records parsed. Check the document format and header definitions.")
        return
    
    # 3. Build pivot DataFrame
    pivot_df = build_pivot_dataframe(records)
    print("Pivot DataFrame:")
    print(pivot_df)
    
    # 4. Create a graph (bar chart for 'Base Value')
    chart_path = create_base_value_chart(pivot_df, graph_image)
    
    # 5. Generate a new DOCX report with the table and graph
    generate_report(pivot_df, chart_path, output_docx)

if __name__ == "__main__":
    input_path = "Data/processed_data/bordertocoast_data/reports/supporting_documents/"
    output_path = "Data/processed_data/bordertocoast_data/reports/supporting_documents/"
    main(input_path, output_path)
