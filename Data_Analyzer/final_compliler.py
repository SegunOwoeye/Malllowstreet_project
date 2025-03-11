import os
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from pathlib import Path

def read_docx_table(docx_path):
    """
    Reads the first table from a DOCX file and returns a DataFrame.
    Assumes the first row is the header.
    """
    doc = Document(docx_path)
    if not doc.tables:
        raise ValueError("No tables found in the document.")
    table = doc.tables[0]
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    data = []
    for row in table.rows[1:]:
        row_data = [cell.text.strip() for cell in row.cells]
        if any(row_data):
            data.append(dict(zip(headers, row_data)))
    return pd.DataFrame(data)

def convert_numeric_columns(df, numeric_cols):
    """
    Convert the specified columns to numeric values (removing commas if needed).
    """
    for col in numeric_cols:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col].str.replace(',', ''), errors='coerce')
    return df

def classify_asset(fund_name):
    """
    Heuristically classifies a fund into an asset class.
    Adjust keywords as needed.
    """
    name = fund_name.lower()
    if "equity" in name or "stocks" in name or "listed" in name:
        return "Public Stocks"
    elif "bond" in name or "credit" in name:
        return "Bonds"
    elif "private-market" in name or "private equity" in name:
        return "Private Equity"
    elif "alternatives" in name:
        return "Alternatives"
    elif "multi-asset" in name:
        return "Multi-Asset"
    else:
        return "Other"

def classify_investment_type(fund_name):
    """
    Classifies a fund as Direct or Indirect.
    Heuristic: if the fund name contains the word 'fund' then it is held indirectly via an investment manager,
    otherwise it is a direct investment.
    """
    if "fund" in fund_name.lower():
        return "Indirect"
    else:
        return "Direct"

def classify_sector(fund_name):
    """
    Heuristically assigns a sector to a fund.
    Since the raw data does not include a sector field,
    this function uses keywords in the fund name.
    You can adjust or extend this mapping as needed.
    """
    name = fund_name.lower()
    # Example mapping (adjust as appropriate):
    if "emerging" in name:
        return "Emerging Markets"
    elif "global" in name:
        return "Global Equities"
    elif "uk" in name:
        return "Public Stocks"
    elif "credit" in name or "bond" in name:
        return "Bonds"
    elif "private-market" in name:
        return "Private Equity"
    else:
        return "Other"

def generate_charts(df, output_dir):
    """
    Generate and save four charts:
      1. Bar chart for asset class breakdown.
      2. Pie chart for direct vs indirect investments.
      3. Bar chart for Base Value by Fund.
      4. Pie chart for sector breakdown.
    Returns a dictionary with paths to the generated images.
    """
    os.makedirs(output_dir, exist_ok=True)
    charts = {}

    # 1. Asset Class Breakdown (Bar Chart)
    asset_class_summary = df.groupby("Asset Class")["Base Value"].sum().reset_index()
    plt.figure(figsize=(8,6))
    plt.bar(asset_class_summary["Asset Class"], asset_class_summary["Base Value"])
    plt.xlabel("Asset Class")
    plt.ylabel("Total Base Value (GBP)")
    plt.title("Asset Class Breakdown")
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()
    asset_class_chart_path = os.path.join(output_dir, "asset_class_breakdown.png")
    plt.savefig(asset_class_chart_path)
    plt.close()
    charts["asset_class"] = asset_class_chart_path

    # 2. Direct vs Indirect Investments (Pie Chart)
    investment_type_summary = df.groupby("Investment Type")["Base Value"].sum().reset_index()
    plt.figure(figsize=(8,6))
    plt.pie(investment_type_summary["Base Value"], labels=investment_type_summary["Investment Type"], autopct='%1.1f%%')
    plt.title("Direct vs Indirect Investments")
    plt.tight_layout()
    investment_type_chart_path = os.path.join(output_dir, "direct_vs_indirect.png")
    plt.savefig(investment_type_chart_path)
    plt.close()
    charts["investment_type"] = investment_type_chart_path

    # 3. Previous Bar Graph: Base Value by Fund
    plt.figure(figsize=(10,6))
    plt.bar(df["Fund"], df["Base Value"])
    plt.xlabel("Fund")
    plt.ylabel("Base Value")
    plt.title("Base Value by Fund")
    plt.xticks(rotation=45, ha="right")
    plt.tight_layout()
    base_value_chart_path = os.path.join(output_dir, "base_value_by_fund.png")
    plt.savefig(base_value_chart_path)
    plt.close()
    charts["base_value_by_fund"] = base_value_chart_path

    # 4. Sector Breakdown (Pie Chart)
    sector_summary = df.groupby("Sector")["Base Value"].sum().reset_index()
    plt.figure(figsize=(8,6))
    plt.pie(sector_summary["Base Value"], labels=sector_summary["Sector"], autopct='%1.1f%%')
    plt.title("Sector Breakdown")
    plt.tight_layout()
    sector_chart_path = os.path.join(output_dir, "sector_breakdown.png")
    plt.savefig(sector_chart_path)
    plt.close()
    charts["sector_breakdown"] = sector_chart_path

    return charts

def generate_report(enhanced_summary, charts, output_docx):
    """
    Generates a DOCX report that includes summary numbers, tables, charts,
    a sector breakdown section, and a clarifying note on investment type definitions.
    """
    doc = Document()
    doc.add_heading("Enhanced Border-to-Coast LGPS Data Report", level=1)
    
    # Overview Summary
    doc.add_heading("Summary Overview", level=2)
    doc.add_paragraph(f"Total Assets Invested: GBP {enhanced_summary['Total Assets']:.2f}")
    doc.add_paragraph(f"UK Investment Exposure: GBP {enhanced_summary['UK Investments']:.2f}")
    
    # Asset Class Breakdown Table & Chart
    doc.add_heading("Asset Class Breakdown", level=2)
    asset_class_data = enhanced_summary['Asset Class Breakdown']
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Asset Class"
    hdr_cells[1].text = "Total Base Value (GBP)"
    for _, row in asset_class_data.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row["Asset Class"])
        row_cells[1].text = f"{row['Base Value']:.2f}"
    doc.add_paragraph("Chart: Asset Class Breakdown")
    doc.add_picture(charts["asset_class"], width=Inches(6))
    
    # Direct vs Indirect Investments Table & Chart
    doc.add_heading("Direct vs Indirect Investments", level=2)
    investment_data = enhanced_summary['Investment Type Breakdown']
    table2 = doc.add_table(rows=1, cols=2)
    hdr_cells = table2.rows[0].cells
    hdr_cells[0].text = "Investment Type"
    hdr_cells[1].text = "Total Base Value (GBP)"
    for _, row in investment_data.iterrows():
        row_cells = table2.add_row().cells
        row_cells[0].text = str(row["Investment Type"])
        row_cells[1].text = f"{row['Base Value']:.2f}"
    doc.add_paragraph("Chart: Direct vs Indirect Investments")
    doc.add_picture(charts["investment_type"], width=Inches(6))
    
    # Previous Bar Graph: Base Value by Fund
    doc.add_heading("Base Value by Fund", level=2)
    doc.add_paragraph("This chart shows the Base Value for each fund:")
    doc.add_picture(charts["base_value_by_fund"], width=Inches(6))
    
    # Sector Breakdown Table & Chart
    doc.add_heading("Sector Breakdown", level=2)
    sector_data = enhanced_summary['Sector Breakdown']
    table3 = doc.add_table(rows=1, cols=2)
    hdr_cells = table3.rows[0].cells
    hdr_cells[0].text = "Sector"
    hdr_cells[1].text = "Total Base Value (GBP)"
    for _, row in sector_data.iterrows():
        row_cells = table3.add_row().cells
        row_cells[0].text = str(row["Sector"])
        row_cells[1].text = f"{row['Base Value']:.2f}"
    doc.add_paragraph("Chart: Sector Breakdown")
    doc.add_picture(charts["sector_breakdown"], width=Inches(6))
    
    
    
    doc.save(output_docx)
    print(f"Enhanced report generated: {output_docx}")

def main(input_path, output_path):
    # Path to the organized report DOCX from your earlier processing
    input_docx = f"{input_path}LGPS_Compiled_Financial_Report_Organized.docx"
    df = read_docx_table(input_docx)
    
    # Convert numeric columns to numbers
    numeric_cols = ["Base Value", "Max of Local Price", "Private-Market-June", "Private-Markets-December", "Shares/Par"]
    df = convert_numeric_columns(df, numeric_cols)
    
    # Calculate total assets invested (sum of Base Value)
    total_assets = df["Base Value"].sum()
    
    # Calculate UK investments: filter funds with "UK" in their name
    uk_mask = df["Fund"].str.contains("UK", case=False, na=False)
    uk_investments = df.loc[uk_mask, "Base Value"].sum()
    
    # Classify funds by asset class, investment type, and add a sector classification
    df["Asset Class"] = df["Fund"].apply(classify_asset)
    df["Investment Type"] = df["Fund"].apply(classify_investment_type)
    df["Sector"] = df["Fund"].apply(classify_sector)
    
    # Group by asset class, investment type, and sector
    asset_class_breakdown = df.groupby("Asset Class")["Base Value"].sum().reset_index()
    investment_type_breakdown = df.groupby("Investment Type")["Base Value"].sum().reset_index()
    sector_breakdown = df.groupby("Sector")["Base Value"].sum().reset_index()
    
    # Prepare enhanced summary data dictionary
    enhanced_summary = {
        "Total Assets": total_assets,
        "UK Investments": uk_investments,
        "Asset Class Breakdown": asset_class_breakdown,
        "Investment Type Breakdown": investment_type_breakdown,
        "Sector Breakdown": sector_breakdown
    }
    
    # Generate charts (including the previous base value chart) and save them to an output folder
    charts = generate_charts(df, output_dir="enhanced_charts")
    
    # Generate the final enhanced DOCX report
    output_docx = f"{output_path}Final_Report.docx"
    path_output = Path(output_docx)
    # Ensure the output directory exists
    path_output.parent.mkdir(parents=True, exist_ok=True)
    generate_report(enhanced_summary, charts, output_docx)
    
if __name__ == "__main__":
    input_path = "Data/processed_data/bordertocoast_data/reports/supporting_documents/"
    output_path = "Data/processed_data/bordertocoast_data/reports/Output/"
    main(input_path, output_path)
